"""
DOCX Report Generator — produces an editable Word document with 10 sections.
"""

from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


PRIMARY_RGB = RGBColor(0x1A, 0x56, 0xA6)
SUCCESS_RGB = RGBColor(0x05, 0x96, 0x69)
WARNING_RGB = RGBColor(0xD9, 0x77, 0x06)
DANGER_RGB = RGBColor(0xDC, 0x26, 0x26)


def _severity_color(severity: str) -> RGBColor:
    s = severity.upper()
    if s == "HIGH":
        return DANGER_RGB
    elif s == "MEDIUM":
        return WARNING_RGB
    return SUCCESS_RGB


def _score_interpretation(score: float) -> str:
    if score > 75:
        return "Strong"
    elif score >= 60:
        return "Good"
    elif score >= 40:
        return "Moderate"
    return "Challenging"


def _add_styled_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY_RGB


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def generate_docx_report(roster: dict, boilerplate_code: str = "") -> bytes:
    """Generate an editable DOCX report from an AgentRoster dict. Returns bytes."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    composite = roster.get("composite_readiness", 0)
    doc_name = roster.get("document_name", "Unknown Document")
    timestamp = roster.get("analysis_timestamp", datetime.now().isoformat())

    # ── 1. COVER PAGE ───────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run("AgentifyX\nTransformation Report")
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY_RGB
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"\nSystem: {doc_name}\nGenerated: {timestamp[:10]}")
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_before = Pt(40)
    run3 = p3.add_run(f"{composite:.0f}/100")
    run3.font.size = Pt(36)
    run3.font.color.rgb = PRIMARY_RGB
    run3.bold = True

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Agentic Readiness Score")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.space_before = Pt(60)
    run_f = footer.add_run("AgentifyX | TECHgium 2026")
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    # ── 2. EXECUTIVE SUMMARY ────────────────────────────────────
    _add_styled_heading(doc, "1. Executive Summary")
    summary = roster.get("transformation_summary", "No summary available.")
    doc.add_paragraph(summary)

    # ── 3. AGENTIC READINESS ASSESSMENT ─────────────────────────
    _add_styled_heading(doc, "2. Agentic Readiness Assessment")
    scores = roster.get("readiness_scores", {})
    dim_names = ["TaskDecomposability", "DecisionComplexity", "ToolIntegrability",
                 "AutonomyPotential", "DataAvailability", "RiskLevel", "ROIPotential"]
    dim_labels = ["Task Decomposability", "Decision Complexity", "Tool Integrability",
                  "Autonomy Potential", "Data Availability", "Risk Level (↑=safer)", "ROI Potential"]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Score"
    hdr[2].text = "Interpretation"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for key, label in zip(dim_names, dim_labels):
        val = scores.get(key, 0)
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{val}/100"
        row[2].text = _score_interpretation(val)

    # Composite row
    comp_row = table.add_row().cells
    comp_row[0].text = "COMPOSITE"
    comp_row[1].text = f"{composite:.0f}/100"
    comp_row[2].text = _score_interpretation(composite)
    for cell in comp_row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()  # spacing

    # ── 4. CURRENT STATE ANALYSIS ───────────────────────────────
    _add_styled_heading(doc, "3. Current State Analysis")
    pain_points = roster.get("legacy_pain_points", [])
    if pain_points:
        for pp in pain_points:
            sev = pp.get("severity", "medium").upper()
            p = doc.add_paragraph()
            run_sev = p.add_run(f"[{sev}] ")
            run_sev.bold = True
            run_sev.font.color.rgb = _severity_color(sev)
            p.add_run(pp.get("pain_point", ""))
    else:
        doc.add_paragraph("No specific pain points identified.")

    # ── 5. PROPOSED AGENT ARCHITECTURE ──────────────────────────
    _add_styled_heading(doc, "4. Proposed Agent Architecture")
    agents = roster.get("agents", [])
    for agent in agents:
        name = agent.get("agent_name", agent.get("role", "Agent"))
        _add_styled_heading(doc, name, level=2)
        doc.add_paragraph(f"Role: {agent.get('role', 'N/A')}")
        doc.add_paragraph(f"Goal: {agent.get('goal', 'N/A')}")
        tools = ", ".join(agent.get("tools_required", []))
        if tools:
            doc.add_paragraph(f"Tools: {tools}")
        hitl = "Yes" if agent.get("hitl_required") else "No"
        doc.add_paragraph(f"HITL Required: {hitl}")
        if agent.get("hitl_required") and agent.get("hitl_checkpoints"):
            doc.add_paragraph(f"Checkpoints: {', '.join(agent['hitl_checkpoints'])}")
        conf = agent.get("confidence_score", 0)
        doc.add_paragraph(f"Confidence: {int(conf*100)}%")
        citations = agent.get("evidence_citations", [])
        for c in citations:
            p = doc.add_paragraph()
            run_cite = p.add_run(f"Source: Page {c.get('page','?')}: {c.get('snippet','')}")
            run_cite.italic = True
            run_cite.font.size = Pt(9)

    # ── 6. HITL REQUIREMENTS ────────────────────────────────────
    _add_styled_heading(doc, "5. HITL Requirements")
    hitl_agents = [a for a in agents if a.get("hitl_required")]
    if hitl_agents:
        ht = doc.add_table(rows=1, cols=3)
        ht.style = "Light Grid Accent 1"
        ht.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = ht.rows[0].cells
        hdr[0].text = "Agent"
        hdr[1].text = "Checkpoint Triggers"
        hdr[2].text = "Risk Level"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for a in hitl_agents:
            row = ht.add_row().cells
            row[0].text = a.get("agent_name", "")
            row[1].text = ", ".join(a.get("hitl_checkpoints", [])) or "General review"
            row[2].text = "Review Required"
    else:
        doc.add_paragraph("No agents require human-in-the-loop approval.")

    # ── 7. FRAMEWORK RECOMMENDATION ─────────────────────────────
    _add_styled_heading(doc, "6. Framework Recommendation")
    fw = roster.get("recommended_framework", "N/A")
    p = doc.add_paragraph()
    p.add_run(f"Recommended: ").bold = False
    run_fw = p.add_run(fw)
    run_fw.bold = True
    run_fw.font.color.rgb = PRIMARY_RGB
    doc.add_paragraph(summary)

    # ── 8. TOOLS RECOMMENDATION ─────────────────────────────────
    _add_styled_heading(doc, "7. Tools Recommendation")
    for agent in agents:
        tools = agent.get("tools_required", [])
        if tools:
            agent_name = agent.get("agent_name", "Agent")
            p = doc.add_paragraph()
            p.add_run(f"{agent_name}: ").bold = True
            p.add_run(", ".join(tools))

    # ── 9. ASSUMPTIONS & DATA GAPS ──────────────────────────────
    _add_styled_heading(doc, "8. Assumptions & Data Gaps")
    assumptions = roster.get("assumptions_made", [])
    gaps = roster.get("data_gaps", [])
    if assumptions:
        _add_styled_heading(doc, "Assumptions Made", level=2)
        for a in assumptions:
            doc.add_paragraph(f"• {a}")
    if gaps:
        _add_styled_heading(doc, "Data Gaps Identified", level=2)
        for g in gaps:
            doc.add_paragraph(f"• {g}")
    if not assumptions and not gaps:
        doc.add_paragraph("None identified.")

    # ── 10. APPENDIX: GENERATED CODE ────────────────────────────
    if boilerplate_code:
        doc.add_page_break()
        _add_styled_heading(doc, "Appendix: Generated Code")
        p_note = doc.add_paragraph()
        run_note = p_note.add_run("Generated by AgentifyX — review before production use")
        run_note.italic = True
        run_note.font.size = Pt(9)

        code_p = doc.add_paragraph()
        code_run = code_p.add_run(boilerplate_code)
        code_run.font.name = "Courier New"
        code_run.font.size = Pt(8)

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
